from __future__ import annotations

import csv
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TABLES = ROOT / "outputs" / "tables"
OUT = ROOT / "outputs" / "manuscript"
OUT.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "183A55"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
WHITE = "FFFFFF"


def read_csv(name: str) -> list[dict[str, str]]:
    with (TABLES / name).open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def font_run(run, size=None, bold=None, italic=None, color=None, name="Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])


def setup_document(title: str, subtitle: str, status: str) -> Document:
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    font_run(header.add_run("Epicenter data analysis suite"), size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font_run(footer.add_run("Page "), size=9, color=MUTED)
    add_page_field(footer)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    font_run(kicker.add_run("EPICENTER ANALYSIS"), size=10, bold=True, color=BLUE)
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(4)
    title_p.paragraph_format.keep_with_next = True
    font_run(title_p.add_run(title), size=24, bold=True, color=NAVY)
    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(12)
    font_run(subtitle_p.add_run(subtitle), size=12.5, color=MUTED)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(16)
    font_run(meta.add_run("Status: "), size=10, bold=True, color=DARK_BLUE)
    font_run(meta.add_run(status), size=10, color=MUTED)
    return doc


def add_callout(doc: Document, heading: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    font_run(p.add_run(heading), size=11, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    font_run(p2.add_run(body), size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_sections(doc: Document, sections: Iterable[dict[str, str]]) -> None:
    for section in sections:
        heading = section["heading"]
        body = section["body"]
        kind = section["kind"]
        level = max(1, min(3, int(section["level"])))
        if kind == "callout":
            add_callout(doc, heading, body)
            continue
        doc.add_heading(heading, level=level)
        if kind == "bullets":
            for line in body.splitlines():
                if line.strip():
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(line.strip())
        elif kind == "code":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.right_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.05
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHT_GRAY)
            p_pr.append(shd)
            font_run(p.add_run(body), size=9.3, name="Consolas")
        else:
            for paragraph in body.split("\n\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())


def add_data_table(doc: Document, title: str, rows: list[dict[str, str]], columns, widths) -> None:
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for index, (_, label) in enumerate(columns):
        cell = header.cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font_run(p.add_run(label), size=8.7, bold=True, color=NAVY)
    for row in rows:
        table_row = table.add_row()
        prevent_row_split(table_row)
        cells = table_row.cells
        for index, (key, _) in enumerate(columns):
            p = cells[index].paragraphs[0]
            if key not in {"label", "analysis", "result", "family", "conclusion"}:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            font_run(p.add_run(str(row.get(key, ""))), size=8.5)
    set_table_geometry(table, widths)
    doc.add_paragraph()


def build_methods_document(sections, item_rows, post_rows, age_rows) -> Path:
    doc = setup_document(
        "Methods and Results Outline",
        "Manuscript-oriented summary of the Epicenter questionnaire analyses",
        "Working scientific outline; exploratory analyses are labelled explicitly",
    )
    add_sections(doc, sections)
    add_data_table(
        doc,
        "Appendix Table A. Item-level PRE-to-POST results",
        item_rows,
        [
            ("item", "Item"), ("label", "Outcome"), ("n_pairs", "N"),
            ("pre_mean_raw", "PRE M"), ("post_mean_raw", "POST M"),
            ("p_fdr", "BH q"), ("rank_biserial_r", "r"), ("conclusion", "Conclusion"),
        ],
        [500, 3300, 600, 700, 700, 700, 650, 2210],
    )
    add_data_table(
        doc,
        "Appendix Table B. Raw POST AE-group comparisons",
        post_rows,
        [
            ("item", "Item"), ("label", "Outcome"), ("n_yes", "N yes"), ("n_no", "N no"),
            ("p_value", "p"), ("p_fdr", "BH q"), ("rank_biserial_r", "r"),
        ],
        [500, 3900, 800, 800, 800, 800, 760],
    )
    add_data_table(
        doc,
        "Appendix Table C. Adjusted age effects on item change",
        age_rows,
        [
            ("item", "Item"), ("label", "Outcome"), ("n", "N"), ("estimate", "Age b"),
            ("conf_low", "CI low"), ("conf_high", "CI high"), ("p_family_fdr", "BH q"),
        ],
        [500, 3700, 650, 850, 850, 850, 960],
    )
    path = OUT / "Epicenter_methods_results_outline.docx"
    doc.save(path)
    return path


def build_audit_document(sections, crosschecks, item_rows, post_rows, ancova_rows, age_rows) -> Path:
    doc = setup_document(
        "Calculation Audit",
        "Manual reconstruction steps, expected values, and pipeline reconciliation",
        "Independent numerical audit; all checks must pass before manuscript submission",
    )
    add_sections(doc, sections)
    add_data_table(
        doc,
        "Audit Table 1. Item-level expected values",
        item_rows,
        [
            ("item", "Item"), ("n_pairs", "N"), ("pre_mean_raw", "PRE M"),
            ("post_mean_raw", "POST M"), ("mean_change_favourable", "Mean change"),
            ("n_improved", "Improved"), ("n_unchanged", "Same"), ("n_worsened", "Worse"),
            ("p_value", "p"), ("p_fdr", "BH q"), ("rank_biserial_r", "r"),
        ],
        [430, 520, 720, 720, 900, 760, 650, 650, 720, 720, 770],
    )
    add_data_table(
        doc,
        "Audit Table 2. Raw POST AE-group tests",
        post_rows,
        [
            ("item", "Item"), ("label", "Outcome"), ("n_yes", "N yes"), ("n_no", "N no"),
            ("p_value", "p"), ("p_fdr", "BH q"), ("rank_biserial_r", "r"),
        ],
        [500, 3900, 800, 800, 800, 800, 760],
    )
    add_data_table(
        doc,
        "Audit Table 3. Baseline-adjusted POST AE coefficients",
        ancova_rows,
        [
            ("item", "Item"), ("label", "Outcome"), ("n", "N"), ("estimate", "AE b"),
            ("conf_low", "CI low"), ("conf_high", "CI high"), ("p_value", "p"), ("p_fdr", "BH q"),
        ],
        [500, 3300, 600, 750, 850, 850, 750, 760],
    )
    add_data_table(
        doc,
        "Audit Table 4. Adjusted age coefficients",
        age_rows,
        [
            ("item", "Item"), ("label", "Outcome"), ("n", "N"), ("estimate", "Age b"),
            ("conf_low", "CI low"), ("conf_high", "CI high"), ("p_value", "p"), ("p_family_fdr", "BH q"),
        ],
        [500, 3300, 600, 750, 850, 850, 750, 760],
    )
    add_data_table(
        doc,
        "Audit Table 5. Independent pipeline reconciliation",
        crosschecks,
        [
            ("family", "Family"), ("result", "Result"), ("metric", "Metric"),
            ("recomputed", "Recomputed"), ("pipeline", "Pipeline"),
            ("absolute_difference", "Abs diff"), ("status", "Status"),
        ],
        [1700, 1300, 1300, 1450, 1450, 1450, 710],
    )
    path = OUT / "Epicenter_calculation_audit.docx"
    doc.save(path)
    return path


def round_fields(rows: list[dict[str, str]], fields: Iterable[str], digits=3) -> list[dict[str, str]]:
    result = []
    for source in rows:
        row = dict(source)
        for field in fields:
            try:
                row[field] = f"{float(row[field]):.{digits}f}"
            except (ValueError, TypeError, KeyError):
                pass
        result.append(row)
    return result


sections = read_csv("manuscript_document_sections.csv")
methods_sections = sorted((x for x in sections if x["document"] == "methods"), key=lambda x: int(x["section_order"]))
audit_sections = sorted((x for x in sections if x["document"] == "audit"), key=lambda x: int(x["section_order"]))
item_rows = round_fields(
    read_csv("manuscript_item_change_summary.csv"),
    ["pre_mean_raw", "post_mean_raw", "mean_change_favourable", "p_value", "p_fdr", "rank_biserial_r"],
)
post_rows = round_fields(read_csv("manuscript_post_ae_tests.csv"), ["p_value", "p_fdr", "rank_biserial_r"])
ancova_rows = round_fields(read_csv("manuscript_post_ancova.csv"), ["estimate", "conf_low", "conf_high", "p_value", "p_fdr"])
age_rows = round_fields(read_csv("manuscript_age_change_models.csv"), ["estimate", "conf_low", "conf_high", "p_value", "p_family_fdr"])
crosschecks = round_fields(read_csv("manuscript_crosscheck.csv"), ["recomputed", "pipeline", "absolute_difference"], digits=10)

methods_path = build_methods_document(methods_sections, item_rows, post_rows, age_rows)
audit_path = build_audit_document(audit_sections, crosschecks, item_rows, post_rows, ancova_rows, age_rows)
print(f"Wrote {methods_path}")
print(f"Wrote {audit_path}")
